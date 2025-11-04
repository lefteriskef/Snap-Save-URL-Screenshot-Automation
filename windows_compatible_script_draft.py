import keyboard
import pyautogui
import pyperclip
from docx import Document
from docx.shared import Inches
import time
import os
from datetime import datetime
import sys
import threading

docxfile = r"C:\Users\lefte\Downloads\Screenshots_project\Snap-Save-URL-Screenshot-Automation\shots.docx"
test = 0
shotfile = r"C:\Users\lefte\Downloads\Screenshots_project\Snap-Save-URL-Screenshot-Automation\shots\shot.png"

if not os.path.exists(docxfile):
    doc = Document()
    doc.save(docxfile)

doc = Document(docxfile)

exit_flag = threading.Event()  # Flag to signal exit

def do_cap():
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    shot = pyautogui.screenshot()
    shot.save(shotfile)
    doc.add_paragraph(f"Screenshot taken at {timestamp}:")
    doc.add_picture(shotfile, width=Inches(7))
    keyboard.press_and_release('ctrl+l')
    time.sleep(1)
    keyboard.press_and_release('ctrl+c')
    time.sleep(1)
    url = pyperclip.paste()
    print(f"--------------Captured URL: {url}")
    doc.add_paragraph(f"Captured URL: {url}")
    doc.save(docxfile)
    print('Screenshot and URL saved and appended to docx.')

def input_listener():
    while not exit_flag.is_set():
        user_input = input()
        if user_input.strip().lower() == "exit":
            print("Exit command received. Exiting script.")
            exit_flag.set()
            keyboard.unhook_all_hotkeys()

keyboard.add_hotkey('ctrl+q', do_cap)

print("Script running. Press Ctrl+Q to take screenshot and grab URL.")
print("Type 'exit' (without quotes) and press Enter to exit safely.")

threading.Thread(target=input_listener, daemon=True).start()

# Replace keyboard.wait() with loop, checking exit_flag
while not exit_flag.is_set():
    time.sleep(0.1)  # Short sleep to reduce CPU use

print("Cleanup done, script exiting.")
sys.exit(0)
